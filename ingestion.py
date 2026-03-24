"""
DocChat — Document Ingestion Module
====================================
Handles document deduplication (SHA-256) and format-aware semantic chunking.
Separates ingestion logic from the UI layer in app.py.
"""

import hashlib
from langchain_core.documents import Document


# ---------------------------------------------------------------------------
# Document Deduplication
# ---------------------------------------------------------------------------

def compute_hash(file_bytes: bytes) -> str:
    """Compute SHA-256 hash of raw file bytes for duplicate detection."""
    return hashlib.sha256(file_bytes).hexdigest()


def is_duplicate(file_hash: str, seen_hashes: dict) -> bool:
    """Check if a document with this hash has already been ingested."""
    return file_hash in seen_hashes


# ---------------------------------------------------------------------------
# Format-Specific Text Extraction
# ---------------------------------------------------------------------------

def extract_text_from_pdf(file) -> list[Document]:
    """Extract text page-by-page from PDF. Each page becomes a separate Document."""
    from pypdf import PdfReader

    docs = []
    reader = PdfReader(file)
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": file.name, "page_number": i + 1, "format": "pdf"}
            ))
    return docs


def extract_text_from_docx(file) -> list[Document]:
    """Extract text from DOCX by heading sections."""
    from docx import Document as DocxDocument

    doc = DocxDocument(file)
    current_section = ""
    current_heading = "Document Start"
    docs = []

    for para in doc.paragraphs:
        # Split on headings to preserve document structure
        if para.style.name.startswith("Heading"):
            if current_section.strip():
                docs.append(Document(
                    page_content=current_section.strip(),
                    metadata={"source": file.name, "section": current_heading, "format": "docx"}
                ))
            current_heading = para.text
            current_section = para.text + "\n"
        else:
            current_section += para.text + "\n"

    # Don't forget the last section
    if current_section.strip():
        docs.append(Document(
            page_content=current_section.strip(),
            metadata={"source": file.name, "section": current_heading, "format": "docx"}
        ))

    return docs


def extract_text_from_csv(file) -> list[Document]:
    """Extract CSV row-by-row. Each row becomes a chunk with column headers as metadata."""
    import csv
    import io

    content = file.read()
    if isinstance(content, bytes):
        content = content.decode("utf-8")

    reader = csv.reader(io.StringIO(content))
    rows = list(reader)

    if len(rows) < 2:
        return []

    headers = rows[0]
    docs = []
    for i, row in enumerate(rows[1:], start=2):
        # Format: "column1: value1 | column2: value2 | ..."
        row_text = " | ".join(f"{h}: {v}" for h, v in zip(headers, row) if v.strip())
        if row_text.strip():
            docs.append(Document(
                page_content=row_text,
                metadata={"source": file.name, "row_number": i, "columns": ", ".join(headers), "format": "csv"}
            ))

    return docs


def extract_text_from_txt(file) -> list[Document]:
    """Extract raw text from TXT files."""
    content = file.read()
    if isinstance(content, bytes):
        content = content.decode("utf-8")

    if content.strip():
        return [Document(
            page_content=content,
            metadata={"source": file.name, "format": "txt"}
        )]
    return []


# ---------------------------------------------------------------------------
# Semantic Chunking
# ---------------------------------------------------------------------------

def semantic_chunk_documents(documents: list[Document], embeddings) -> list[Document]:
    """
    Apply semantic chunking to extracted documents.
    - CSV rows are NOT chunked further (each row is already a chunk).
    - Short documents (< 100 chars) are kept as-is.
    - Everything else gets semantic chunking based on topic boundaries.
    """
    from langchain_experimental.text_splitter import SemanticChunker

    chunker = SemanticChunker(
        embeddings,
        breakpoint_threshold_type="percentile",
        breakpoint_threshold_amount=95,
    )

    chunked_docs = []

    for doc in documents:
        # CSV rows: already atomic, don't chunk further
        if doc.metadata.get("format") == "csv":
            chunked_docs.append(doc)
            continue

        # Very short content: not worth chunking
        if len(doc.page_content) < 100:
            chunked_docs.append(doc)
            continue

        # Semantic chunking
        try:
            chunks = chunker.create_documents(
                [doc.page_content],
                metadatas=[doc.metadata]
            )
            chunked_docs.extend(chunks)
        except Exception:
            # Fallback: if semantic chunking fails, keep original
            chunked_docs.append(doc)

    return chunked_docs


# ---------------------------------------------------------------------------
# Main Processing Pipeline
# ---------------------------------------------------------------------------

def process_documents(uploaded_files, seen_hashes: dict, embeddings) -> tuple[list[Document], dict, list[str], list[str]]:
    """
    Full ingestion pipeline:
    1. Hash check → skip duplicates
    2. Extract text per format
    3. Semantic chunking
    4. Return (documents, updated_hashes, processed_names, skipped_names)
    """
    all_docs = []
    processed_names = []
    skipped_names = []

    for file in uploaded_files:
        # Read raw bytes for hashing
        file_bytes = file.read()
        file.seek(0)  # Reset for text extraction

        file_hash = compute_hash(file_bytes)

        if is_duplicate(file_hash, seen_hashes):
            skipped_names.append(file.name)
            continue

        # Extract text based on format
        if file.name.endswith(".pdf"):
            docs = extract_text_from_pdf(file)
        elif file.name.endswith(".docx"):
            docs = extract_text_from_docx(file)
        elif file.name.endswith(".csv"):
            docs = extract_text_from_csv(file)
        elif file.name.endswith(".txt"):
            docs = extract_text_from_txt(file)
        else:
            continue

        # Add doc_hash to metadata for tracking
        for doc in docs:
            doc.metadata["doc_hash"] = file_hash

        all_docs.extend(docs)
        processed_names.append(file.name)

        # Register hash
        seen_hashes[file_hash] = {
            "filename": file.name,
            "chunk_count": len(docs),
        }

    # Semantic chunking on all extracted documents
    if all_docs:
        all_docs = semantic_chunk_documents(all_docs, embeddings)

    return all_docs, seen_hashes, processed_names, skipped_names
